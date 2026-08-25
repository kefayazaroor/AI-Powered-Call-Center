from flask import (
    Flask, render_template, request, jsonify,
    redirect, url_for, session, flash, send_file
)
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import whisper
import os
from datetime import datetime, date
from functools import wraps
from io import BytesIO
import secrets
from sqlalchemy import func
from docx import Document
from docx.shared import Pt, Inches
import qrcode
from flask import g
import json, os 


translations = {
    "en": {
        "login_title": "Login - Call Center System",
        "login_header": "AI Call Center",
        "login_subtitle": "AI system for analyzing call center conversations",
        "login_role_tag": "Employee / Admin Login",
        "username_label": "Email / Username",
        "password_label": "Password",
        "role_label": "Login Type",
        "role_agent": "Call Center Agent",
        "role_admin": "System Admin",
        "login_btn": "Login",
        "demo_data": "Demo accounts",
        "demo_admin": "Admin: admin / 1234",
        "demo_agent": "Agent: agent / 1234",
    },

    "ar": {}  # العربي فاضي لأنه النصوص الأصلية موجودة
}



# ===================== تهيئة التطبيق =====================

app = Flask(__name__)
app.secret_key = "super_secret_key_change_me"
app.config['SECRET_KEY'] = 'change_this_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///callcenter.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# ===================== نماذج قاعدة البيانات =====================

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), nullable=False, default="agent")
    is_active = db.Column(db.Boolean, default=True)
    api_key = db.Column(db.String(64), unique=True, nullable=True)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)


class Call(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    call_id = db.Column(db.String(50), unique=True, nullable=False)

    employee_name = db.Column(db.String(120))
    phone = db.Column(db.String(50))
    case_type = db.Column(db.String(50))

    date = db.Column(db.String(20))
    time = db.Column(db.String(20))

    transcript = db.Column(db.Text)
    summary = db.Column(db.Text)

    sentiment = db.Column(db.String(20))
    urgency = db.Column(db.String(20))

    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    user_id = db.Column(db.Integer, db.ForeignKey("user.id"))
    user = db.relationship("User", backref="calls")


# ===================== تهيئة قاعدة البيانات =====================

def generate_api_key():
    return secrets.token_hex(16)

with app.app_context():
    db.create_all()

    admin = User.query.filter_by(username="admin").first()
    if not admin:
        admin_user = User(
            username="admin",
            password_hash=generate_password_hash("admin123"),
            role="admin",
            is_active=True,
            api_key=generate_api_key()
        )
        db.session.add(admin_user)
        db.session.commit()
        print("Created default admin user: admin / admin123")

# ===================== تحميل Whisper =====================

model = whisper.load_model("small")  # tiny لو بدك بدليه

# ===================== دوال المساعدة =====================

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        if session.get("role") != "admin":
            flash("هذه الصفحة خاصة بالمدير.", "error")
            return redirect(url_for("index"))
        return f(*args, **kwargs)
    return decorated


def analyze_call(employee_name, phone, case_type, audio_file, user_id=None):

    os.makedirs("uploads", exist_ok=True)
    filename = f"call_{datetime.now().strftime('%Y%m%d_%H%M%S')}.wav"
    filepath = os.path.join("uploads", filename)
    audio_file.save(filepath)

    result = model.transcribe(filepath, language="ar")
    transcript = (result.get("text") or "").strip()

    summary = transcript
    if len(summary) > 400:
        summary = summary[:400] + "..."

    text_lower = transcript

    sentiment = "محايد"
    negative_words = ["غلط", "سيء", "سيئة", "تعبت", "مش راضي", "مش راضية", "مشكلة", "شكوى", "زعلان"]
    positive_words = ["شكراً", "شكرا", "ممتاز", "كويس", "راضي", "راضية", "تمام"]

    if any(w in text_lower for w in negative_words):
        sentiment = "سلبي"
    if any(w in text_lower for w in positive_words):
        sentiment = "إيجابي"

    urgency = "عادي"
    urgent_words = ["ضروري", "مستعجل", "اليوم", "هلأ", "هسا", "حالاً", "طوارئ"]
    if any(w in text_lower for w in urgent_words):
        urgency = "عالي"

    now = datetime.now()
    date_str = now.strftime("%Y-%m-%d")
    time_str = now.strftime("%H:%M")

    detected_case_type = case_type or "غير محدد"

    call = Call(
        call_id=filename.replace(".wav", ""),
        employee_name=employee_name,
        phone=phone,
        case_type=detected_case_type,
        date=date_str,
        time=time_str,
        transcript=transcript,
        summary=summary,
        sentiment=sentiment,
        urgency=urgency,
        user_id=user_id,
    )
    db.session.add(call)
    db.session.commit()

    return {
        "employee_name": employee_name,
        "phone": phone,
        "case_type": detected_case_type,
        "date": date_str,
        "time": time_str,
        "transcript": transcript,
        "summary": summary,
        "sentiment": sentiment,
        "urgency": urgency,
        "call_id": call.call_id,
        "db_id": call.id,
    }


# ===================== الدخول والخروج =====================

def load_translation(lang):
    path = f"translations/{lang}.json"
    if os.path.exists(path):
        return json.load(open(path, "r", encoding="utf-8"))
    return {}

@app.before_request
def set_language():
    lang = session.get("lang", "ar")
    g.lang = lang
    g.tr = load_translation(lang)

def tr(key, fallback=None):
    return g.tr.get(key, fallback or key)
 

@app.route("/set-lang/<lang>")
def set_lang(lang):
    if lang not in ["ar", "en"]:
        lang = "ar"
    session["lang"] = lang
    return redirect(request.referrer or url_for("login"))



@app.before_request
def before_request():
    g.lang = session.get("lang", "ar")

app.jinja_env.globals.update(tr=tr)


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        user = User.query.filter_by(username=username).first()
        if user and user.is_active and user.check_password(password):
            session["user_id"] = user.id
            session["username"] = user.username
            session["role"] = user.role

            if user.role == "admin":
                return redirect(url_for("dashboard"))
            else:
                return redirect(url_for("index"))

        flash("بيانات الدخول غير صحيحة أو الحساب غير مفعل.", "error")

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ===================== صفحات الموظف =====================

@app.route("/landing")
def landing():
    return render_template("landing.html")


@app.route("/")
def home():
    return redirect(url_for("landing"))


@app.route("/index")
@login_required
def index():
    return render_template(
        "index.html",
        username=session.get("username"),
        role=session.get("role")
    )


# ===================== API داخلي =====================

@app.route("/api/transcribe", methods=["POST"])
@login_required
def api_transcribe():
    employee_name = request.form.get("employee_name", "").strip()
    phone = request.form.get("phone", "").strip()
    case_type = request.form.get("case_type", "").strip()
    audio_file = request.files.get("audio")

    if not audio_file:
        return jsonify({"error": "لم يتم إرسال ملف صوت"}), 400

    data = analyze_call(employee_name, phone, case_type, audio_file, user_id=session["user_id"])
    return jsonify(data)


# ===================== API خارجي =====================

@app.route("/api/v1/transcribe", methods=["POST"])
def external_api_transcribe():

    api_key = request.headers.get("X-API-KEY", "").strip()
    if not api_key:
        return jsonify({"error": "مطلوب X-API-KEY"}), 401

    user = User.query.filter_by(api_key=api_key, is_active=True).first()
    if not user:
        return jsonify({"error": "API KEY غير صالح"}), 401

    employee_name = request.form.get("employee_name", "").strip()
    phone = request.form.get("phone", "").strip()
    case_type = request.form.get("case_type", "").strip()
    audio_file = request.files.get("audio")

    if not audio_file:
        return jsonify({"error": "لم يتم إرسال ملف صوت"}), 400

    data = analyze_call(employee_name, phone, case_type, audio_file, user_id=user.id)
    return jsonify(data)


# ===================== Dashboard =====================

@app.route("/dashboard")
@admin_required
def dashboard():
    employees_count = User.query.count()

    total_calls = Call.query.count()
    today_str = date.today().strftime("%Y-%m-%d")
    todays_calls = Call.query.filter_by(date=today_str).count()

    complaints = Call.query.filter_by(case_type="شكوى").count()
    inquiries = Call.query.filter_by(case_type="استفسار").count()
    followups = Call.query.filter_by(case_type="متابعة").count()

    positive = Call.query.filter_by(sentiment="إيجابي").count()
    negative = Call.query.filter_by(sentiment="سلبي").count()
    neutral = Call.query.filter_by(sentiment="محايد").count()

    urgent = Call.query.filter_by(urgency="عالي").count()
    normal = Call.query.filter_by(urgency="عادي").count()

    recent_calls = Call.query.order_by(Call.created_at.desc()).limit(20).all()

    stats_by_date = (
        db.session.query(Call.date, func.count(Call.id))
        .group_by(Call.date)
        .order_by(Call.date)
        .all()
    )

    dates = [row[0] for row in stats_by_date]
    counts = [row[1] for row in stats_by_date]

    return render_template(
        "admin_dashboard.html",
        total_calls=total_calls,
        todays_calls=todays_calls,
        complaints=complaints,
        inquiries=inquiries,
        followups=followups,
        positive=positive,
        negative=negative,
        neutral=neutral,
        urgent=urgent,
        normal=normal,
        recent_calls=recent_calls,
        dates=dates,
        counts=counts,
        employees_count=employees_count,
        username=session.get("username")
    )

# ===================== إدارة المستخدمين (مدير فقط) =====================

@app.route("/admin/users")
@admin_required
def manage_users():
    users = User.query.order_by(User.id.asc()).all()
    return render_template(
        "users.html",
        users=users,
        username=session.get("username")
    )


@app.route("/admin/users/create", methods=["GET", "POST"])
@admin_required
def create_user():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        if not username or not password:
            flash("يرجى إدخال اسم مستخدم وكلمة مرور.", "error")
            return redirect(url_for("create_user"))

        if User.query.filter_by(username=username).first():
            flash("اسم المستخدم مستخدم مسبقاً", "error")
            return redirect(url_for("create_user"))

        new_user = User(
            username=username,
            password_hash=generate_password_hash(password),
            role="agent",
            is_active=True,
            api_key=generate_api_key()
        )

        db.session.add(new_user)
        db.session.commit()

        flash("تم إنشاء المستخدم بنجاح.", "success")
        return redirect(url_for("manage_users"))

    return render_template("create_user.html", username=session.get("username"))


@app.route("/admin/users/<int:user_id>/edit", methods=["GET", "POST"])
@admin_required
def edit_user(user_id):
    user = User.query.get_or_404(user_id)

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        if username:
            user.username = username

        if password:
            user.password_hash = generate_password_hash(password)

        db.session.commit()
        flash("تم تحديث بيانات المستخدم.", "success")
        return redirect(url_for("manage_users"))

    return render_template("edit_user.html", user=user, username=session.get("username"))


@app.route("/admin/users/<int:user_id>/toggle-active", methods=["POST"])
@admin_required
def toggle_user_active(user_id):
    user = User.query.get_or_404(user_id)
    user.is_active = not user.is_active
    db.session.commit()
    flash("تم تغيير حالة الحساب.", "success")
    return redirect(url_for("manage_users"))


@app.route("/admin/users/<int:user_id>/reset-api-key", methods=["POST"])
@admin_required
def reset_api_key(user_id):
    user = User.query.get_or_404(user_id)
    user.api_key = generate_api_key()
    db.session.commit()
    flash("تم تجديد API KEY", "success")
    return redirect(url_for("manage_users"))


@app.route("/admin/users/<int:user_id>/delete", methods=["POST"])
@admin_required
def delete_user(user_id):
    user = User.query.get_or_404(user_id)

    if user.username == "admin":
        flash("لا يمكن حذف مدير النظام!", "error")
        return redirect(url_for("manage_users"))

    db.session.delete(user)
    db.session.commit()
    flash("تم حذف المستخدم.", "success")
    return redirect(url_for("manage_users"))

# ===================== صفحة جميع المكالمات (إصدار نهائي) =====================

@app.route("/all-calls")
@admin_required
def all_calls():

    search = request.args.get("search", "").strip()
    case_filter = request.args.get("case_type", "").strip()
    date_from = request.args.get("date_from", "").strip()
    date_to = request.args.get("date_to", "").strip()

    query = Call.query.order_by(Call.created_at.desc())

    if search:
        s = f"%{search}%"
        query = query.filter(
            (Call.employee_name.like(s)) |
            (Call.phone.like(s)) |
            (Call.case_type.like(s)) |
            (Call.date.like(s))
        )

    if case_filter:
        query = query.filter(Call.case_type == case_filter)

    if date_from:
        query = query.filter(Call.date >= date_from)

    if date_to:
        query = query.filter(Call.date <= date_to)

    calls = query.all()

    return render_template(
        "all_calls.html",
        calls=calls,
        search=search,
        case_filter=case_filter,
        date_from=date_from,
        date_to=date_to,
        username=session.get("username")
    )

# ===================== تفاصيل مكالمة =====================

@app.route("/call/<int:call_id>")
@login_required
def call_detail(call_id):
    call = Call.query.get_or_404(call_id)
    return render_template(
        "call_detail.html",
        call=call,
        username=session.get("username"),
        role=session.get("role")
    )

# ===================== API DOCS =====================

@app.route("/api/docs")
@admin_required
def api_docs():
    return render_template("api_docs.html", username=session.get("username"))

# ===================== PDF =====================

@app.route("/call/<int:call_id>/pdf")
@login_required
def call_pdf(call_id):
    call = Call.query.get_or_404(call_id)

    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import mm

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = height - 20 * mm
    c.setFont("Helvetica-Bold", 14)
    c.drawString(20 * mm, y, "تقرير مكالمة كول سنتر (AI Call Center Report)")
    y -= 10 * mm

    c.setFont("Helvetica", 10)
    lines = [
        f"معرّف المكالمة: {call.call_id}",
        f"اسم الموظف: {call.employee_name}",
        f"رقم العميل: {call.phone}",
        f"نوع الحالة: {call.case_type}",
        f"التاريخ: {call.date}  الوقت: {call.time}",
        f"حالة العميل: {call.sentiment}",
        f"درجة الأهمية: {call.urgency}",
        "",
        "ملخص المكالمة:",
        call.summary or "",
        "",
        "النص الكامل:",
        call.transcript or "",
    ]

    for line in lines:
        for chunk in [line[i:i+90] for i in range(0, len(line), 90)]:
            if y < 20 * mm:
                c.showPage()
                y = height - 20 * mm
                c.setFont("Helvetica", 10)
            c.drawString(20 * mm, y, chunk)
            y -= 6 * mm

    c.showPage()
    c.save()
    buffer.seek(0)

    filename = f"call_report_{call.call_id}.pdf"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/pdf"
    )

# ===================== WORD =====================

@app.route("/call/<int:call_id>/word")
@login_required
def call_word(call_id):
    call = Call.query.get_or_404(call_id)

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    document.add_heading("تقرير مكالمة - نظام الكول سنتر الذكي", level=1)

    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    def add_row(label, value):
        row = table.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)

    add_row("معرّف المكالمة", call.call_id)
    add_row("اسم الموظف", call.employee_name)
    add_row("رقم العميل", call.phone)
    add_row("نوع الحالة", call.case_type)
    add_row("التاريخ", call.date)
    add_row("الوقت", call.time)
    add_row("المشاعر", call.sentiment)
    add_row("درجة الأهمية", call.urgency)

    document.add_paragraph("\n")

    document.add_heading("ملخص المكالمة", level=2)
    document.add_paragraph(call.summary or "")

    document.add_heading("النص الكامل للمكالمة", level=2)
    document.add_paragraph(call.transcript or "")

    document.add_paragraph("\n")

    qr_data = f"http://127.0.0.1:5000/call/{call.id}"
    qr_img = qrcode.make(qr_data)
    qr_path = f"qr_{call.call_id}.png"
    qr_img.save(qr_path)

    document.add_heading("رمز QR لفتح المكالمة", level=2)
    document.add_picture(qr_path, width=Inches(2))

    if os.path.exists(qr_path):
        os.remove(qr_path)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"call_report_{call.call_id}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )




if __name__ == "__main__":
    app.run(debug=True)

